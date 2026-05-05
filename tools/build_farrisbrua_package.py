from pathlib import Path
import json
import re
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "produksjon_farrisbrua"
IMG_DIR = OUT / "08_bildepakke"
QA = OUT / "qa"


SCENARIO = {
    "title": "Farrisbrua B43",
    "subtitle": "Trafikkulykke med militær kolonne og sekundær brann/røykutvikling",
    "sted": "E18 Farrisbrua, Larvik, fiktivt B43-område",
    "retning": "Retning Oslo/Sandefjord",
    "sentral": "B43 fiktiv øvelsessentral",
    "kurssted": "Stavern",
    "ovede": "4 personer: 3 operatører og 1 vaktleder",
    "action_time": "ca. 90 minutter",
    "form": "Spilløvelse i kontrollrom",
    "dato": "Utkast 04.05.2026",
}


OBJECTIVES = [
    {
        "id": "D1",
        "mål": "Vaktlaget skal håndtere første nødmelding og utalarmering strukturert og tidsriktig.",
        "indikator": "Ressurser og trippelvarsling iverksettes på tidlig og forsvarlig grunnlag, uten å vente på perfekt informasjon.",
    },
    {
        "id": "D2",
        "mål": "Vaktlaget skal etablere og oppdatere felles situasjonsforståelse mens ressurser er på vei.",
        "indikator": "VL og operatører deler nøkkelinformasjon om posisjon, skade, brann/røyk, trafikk og usikkerhet.",
    },
    {
        "id": "D3",
        "mål": "Vaktlaget skal samvirke hensiktsmessig med politi, AMK og VTS i en hendelse med totalforsvarskontekst.",
        "indikator": "Aktørene kobles inn med presis informasjon og oppdateres ved vesentlige endringer.",
    },
    {
        "id": "D4",
        "mål": "Vaktleder skal fordele oppgaver og sikre overgang fra akuttfase til driftsfase.",
        "indikator": "Det etableres tydelig rollefordeling tidlig, og vaktlaget skifter arbeidsform når første ressurs er fremme.",
    },
    {
        "id": "D5",
        "mål": "Vaktlaget skal oppsummere status presist før øvelsen avsluttes.",
        "indikator": "Status inneholder hva som har skjedd, hva som er gjort, hvem som er varslet, og hva som fortsatt følges opp.",
    },
]


CALLERS = [
    {
        "nr": "1",
        "tid": "00:00",
        "rolle": "Første melder, bilfører rett bak ulykken",
        "navn": "Marius Eik",
        "spilles_av": "Vibecke",
        "kanal": "Nødtelefon 110",
        "uttrykk": "Stresset, men samarbeidsvillig. Puster raskt, snakker fort og avbryter seg selv.",
        "førstemelding": "Det har krasjet på Farrisbrua på E18 ved Larvik, retning Oslo. Det ryker fra en personbil som har kjørt inn i et militært kjøretøy. Jeg tror det er folk skadet.",
        "frivillig": [
            "Han står 30-50 meter bak ulykken.",
            "Trafikken har stoppet i begge felt i samme retning.",
            "Han ser ett sivilt kjøretøy med frontskade og røyk fra motorrom.",
            "Han ser 2-3 personer i uniform ute ved et større kjøretøy.",
        ],
        "ved_spørsmål": [
            "En person sitter i personbilen og virker våken, men kommer seg ikke ut.",
            "En person står utenfor og holder seg til armen/skulderen.",
            "Det lukter diesel, men han ser ikke tydelig væske renne.",
            "Det er ikke synlig åpen flamme ved første samtale.",
            "Han er usikker på nøyaktig retning, men sier 'mot Oslo/Sandefjord' hvis operatør spør.",
        ],
        "usikkerhet": "Kan først si 'det brenner' selv om han egentlig ser røyk. Må korrigeres gjennom spørsmål.",
        "avslutning": "Holdes i samtale så lenge det er nyttig for posisjon, skadeomfang og sikkerhet. Kan brukes som ressurs til å holde avstand og varsle andre bilister om ikke å gå tett på.",
        "forventet": "Trippelvarsling og utalarmering skal starte på denne meldingen.",
    },
    {
        "nr": "2",
        "tid": "03:30",
        "rolle": "Melder fra avstand, ser røyk fra nærområdet",
        "navn": "Anne Løwe",
        "spilles_av": "Benjamin",
        "kanal": "Nødtelefon 110",
        "uttrykk": "Rolig, bekymret, har begrenset informasjon.",
        "førstemelding": "Jeg ser røyk oppe ved Farrisbrua. Det står mye biler stille. Jeg vet ikke hva som har skjedd.",
        "frivillig": [
            "Ser hendelsen fra avstand.",
            "Kan ikke bekrefte skadde eller brann.",
            "Ser kø og noe røyk i retning brua.",
        ],
        "ved_spørsmål": [
            "Ingen nye opplysninger om personer i fare.",
            "Ingen direkte observasjon av ulykkesstedet.",
        ],
        "usikkerhet": "Kan overdrive omfang fordi hun ser kø og røyk, men har ikke nærkontakt.",
        "avslutning": "Avsluttes kort: melding er mottatt, nødetater er på vei, ring tilbake ved ny tidskritisk informasjon.",
        "forventet": "Operatør skal sjekke om melder har ny kritisk informasjon og avslutte kort uten å binde kapasitet.",
    },
    {
        "nr": "3",
        "tid": "05:30",
        "rolle": "Melder i kø fra motsatt/tilstøtende kjørefelt",
        "navn": "Sivert Holm",
        "spilles_av": "Benjamin",
        "kanal": "Nødtelefon 110",
        "uttrykk": "Utålmodig og litt irritert, men ikke panisk.",
        "førstemelding": "Det er full stopp ved Farrisbrua. Jeg ser blålys langt fremme og noen militære biler. Er dette meldt?",
        "frivillig": [
            "Ser hendelsen på avstand.",
            "Oppfatter at militære kjøretøy er involvert.",
            "Har ikke sett brann eller skadde selv.",
        ],
        "ved_spørsmål": [
            "Ingen nye tidskritiske opplysninger.",
            "Han står trygt i kø og er ikke involvert.",
        ],
        "usikkerhet": "Sier 'militærulykke' uten egentlig å vite hva som har skjedd.",
        "avslutning": "Avsluttes kort: melding er mottatt, nødetater er på vei, følg anvisning fra politi/VTS.",
        "forventet": "Operatør skal avgrense samtalen og ikke la totalforsvarskonteksten trekke fokus bort fra hendelsen.",
    },
]


RESOURCES = [
    ["B43 OP 1", "Øvende", "Mottak første nødtelefon, melderoppfølging, livreddende veiledning ved behov"],
    ["B43 OP 2", "Øvende", "Utalarmering, trippelvarsling, AMK/politi og VTS-samvirke"],
    ["B43 OP 3", "Øvende", "Logg, ressursoversikt, støtte til VL og oppfølging av 01/09/UL/IL"],
    ["B43 VL", "Øvende", "Rollefordeling, prioritering, felles situasjonsbilde og overgang til driftsfase"],
    ["B43-01", "Spilt brannressurs", "Første ressurs underveis og fremme"],
    ["B43-09 / UL", "Spilt brannressurs", "Underveismelding, vindusmelding og taktiske avklaringer"],
    ["IL brann", "Spilt brannressurs", "Driftsfase, status og eventuelle ressursbehov"],
    ["AMK", "Motspill", "Trippelvarsling, skadde, ambulanseressurs og pasientprioritering"],
    ["Politi", "Motspill", "Trippelvarsling, trafikk, sikring, militær kolonne som kontekst"],
    ["VTS", "Motspill", "Trafikkavvikling, stenging, kø og fremkommelighet"],
    ["Militær kolonneleder", "Begrenset motspill", "Enkel faktakilde, kan spilles via politi/VTS hvis egen markør mangler"],
]


ROLE_ALLOCATION = [
    ["Rune", "Gruppe 3", "Spillstab", "Øvingsleder / spill-leder", "Start/stopp, tempo, avvik fra dreiebok og overgang til AAR."],
    ["Mats", "Gruppe 3", "Spillstab", "Brannressurser / teknisk spill", "Spiller B43-01, B43-09/UL og IL brann etter dreiebok."],
    ["Siw", "Gruppe 3", "Kontrollrom", "Observatør 1", "D1/D3: mottak, utalarmering, trippelvarsling og samvirke."],
    ["Tom Christian", "Gruppe 3", "Kontrollrom", "Observatør 2", "D2/D4/D5: situasjonsbilde, rollefordeling, driftsfase og sluttstatus."],
    ["Vibecke", "Gruppe 1", "Spillstab", "Første melder / bildestøtte", "Spiller innringer 1 og støtter bilde 1-2 etter avtale."],
    ["Benjamin", "Gruppe 1", "Spillstab", "Andre/tredje melder + VTS", "Spiller innringer 2 og 3, samt VTS ved trafikkavklaringer."],
    ["Johannes", "Gruppe 1", "Spillstab", "AMK", "Spiller AMK i trippel og skade-/pasientavklaringer."],
    ["Ruben", "Gruppe 1", "Spillstab", "Politi / militær kolonneleder", "Spiller politi og enkel kolonnelederavklaring via politi."],
]


TIMELINE = [
    ["00:00", "B43 OP", "Melder 1", "Nødtelefon: trafikkulykke på E18 Farrisbrua retning Oslo/Sandefjord. Røyk fra personbil, mulig skadde, militært kjøretøy involvert.", "D1", "Motta, posisjonere, vurdere brann/skade/risiko, starte utalarmering og trippelvarsling.", "Første melding skal være grunnlag for tiltak."],
    ["01:10", "B43 / ressurser", "Øvingsledelse", "Forventet tidspunkt: brannressurser varslet eller under utalarmering.", "D1", "Ikke vente på fullstendig informasjon før respons.", "70-110 sekunder brukes som realistisk kvalitetsmål, ikke stoppeklokke-felle."],
    ["02:00", "AMK/politi", "B43 OP/VL", "Trippelvarsling etablert. Del posisjon, hendelsestype, skadde, røyk/brann og trafikale forhold.", "D1/D3", "Gi presis og usikkerhetsbevisst informasjon.", "AMK og politi stiller korte oppfølgingsspørsmål."],
    ["03:30", "B43 OP", "Melder 2", "Melder ser røyk og kø fra avstand.", "D2", "Avklare om det er ny tidskritisk informasjon. Avslutte kort.", "Skal ikke binde unødig kapasitet."],
    ["05:30", "B43 OP", "Melder 3", "Melder i kø ser militære biler og spør om det er meldt.", "D2/D4", "Avklare om melder har ny info. Avslutte kort og berolige.", "Totalforsvarskontekst skal ikke overta fokus."],
    ["08:00", "B43 OP/VL", "VTS", "VTS kan bekrefte rask køoppbygging og spør om behov for stenging/varsling.", "D3", "Samvirke om trafikk og fremkommelighet.", "VTS skal være støtteaktør, ikke hovedhendelse."],
    ["10:00", "B43 OP/VL", "B43-01/09 underveis", "Første brannressurs ber om siste status: posisjon, røyk/flamme, skadde, trafikk og eventuell farlig last.", "D2", "Gi kort felles situasjonsbilde til ressurs underveis.", "Tester informasjonsdeling under fremkjøring."],
    ["15:00", "B43 OP", "AMK", "AMK ber om oppdatert antall skadde og om noen sitter fast.", "D2/D3", "Dele bekreftet og ubekreftet informasjon tydelig.", "Unngå å gjøre antakelser til fakta."],
    ["18:00", "B43 OP", "Politi", "Politiet spør om militær kolonne har farlig last eller ammunisjon.", "D3", "Formidle hva som er kjent/ukjent. Ikke spekulere.", "Kan besvares med at dette må avklares via politi/kolonneleder."],
    ["22:00", "B43 OP/VL", "B43-09 / UL", "Vindusmelding: Trafikkulykke personbil mot militært logistikkjøretøy. Røyk fra motorrom, ikke åpen flamme. En person fastklemt/innesperret, en lettere skadd ute. Dieselukt, ikke bekreftet lekkasje. E18 retning Oslo blokkert.", "D2/D4", "Motta, loggføre, dele med AMK/politi/VTS og gå mot driftsfase.", "Vendepunkt i øvelsen."],
    ["30:00", "B43 OP", "UL", "UL ønsker VTS/politi oppdatert på behov for full stans i retning Oslo og trygg arbeidsplass.", "D3/D4", "Koordinere uten å overta skadestedsledelse.", "110 støtter taktisk nivå."],
    ["35:00", "B43 OP/VL", "IL brann", "IL fremme. Ber om samlet status fra sentralen: varslede aktører, meldinger, skadde, trafikk og uavklarte forhold.", "D4/D5", "Gi samlet situasjonsbilde og avklare videre oppfølging.", "Overgang til driftsfase."],
    ["45:00", "B43 OP", "VTS", "VTS melder økende kø og spør om forventet varighet/kommunikasjonslinje videre.", "D3/D4", "Avklare med IL/politi og gi realistisk status.", "Driftsfaseinnspill."],
    ["55:00", "B43 OP", "Politi", "Politiet bekrefter at militær kolonneleder er ivaretatt på stedet. Spør om 110 har mottatt flere meldinger om brann/røyk.", "D3/D4", "Oppsummere meldingsbildet kort.", "Forsvaret holdes som enkel kontekst."],
    ["65:00", "B43 OP", "AMK", "AMK melder at pasient håndteres og ber om oppdatering dersom brannfare endres.", "D3/D4", "Loggføre og sikre intern deling.", "Lavintensiv driftsfase."],
    ["75:00", "B43 OP/VL", "IL brann", "IL ber om oppdatert loggstatus og hvilke aktører som er varslet/kontaktet.", "D5", "VL/OP gir strukturert status.", "Forberedelse til avslutning."],
    ["85:00", "Alle", "Øvingsleder", "Øvelsen fryses. Vaktlaget bes gi kort sluttstatus før AAR.", "D5", "Hva har skjedd, hva er gjort, hvem er varslet, hva gjenstår?", "Avslutning."],
]


SPILLCARDS = [
    {
        "rolle": "Første melder",
        "spiller": "Vibecke",
        "mål": "Starte hendelsen og gi realistisk, ufullstendig, men tilstrekkelig informasjon til tidlig respons.",
        "vet": "Ulykken er på Farrisbrua, retning Oslo/Sandefjord. Personbil med røyk fra motorrom. Militært kjøretøy involvert. Minst en person er skadd/innesperret.",
        "vet_ikke": "Eksakt farlig last, eksakt skadegrad, om det er reell brann eller bare røyk.",
        "ikke_impro": "Ikke oppgi eksplosiver, ammunisjon, kjemikalier eller mange hardt skadde uten at øvingsleder beslutter det.",
        "forventet": "Operatør styrer samtalen, innhenter posisjon og tidskritisk informasjon, starter tiltak og gir trygg veiledning.",
    },
    {
        "rolle": "Andre melder",
        "spiller": "Benjamin",
        "mål": "Teste kort håndtering av duplikatmelding fra avstand.",
        "vet": "Ser røyk og kø fra avstand.",
        "vet_ikke": "Skadde, brannomfang, eksakt kjøretøytype.",
        "ikke_impro": "Ikke legg til ny dramatikk.",
        "forventet": "Operatør avklarer om det er ny tidskritisk informasjon og avslutter kort.",
    },
    {
        "rolle": "Tredje melder",
        "spiller": "Benjamin",
        "mål": "Teste at totalforsvarsord som 'militær' ikke skaper unødvendig sidespor.",
        "vet": "Ser militære kjøretøy og kø fra avstand.",
        "vet_ikke": "Hva som faktisk har skjedd.",
        "ikke_impro": "Ikke hevde at det er terror, sabotasje eller farlig last.",
        "forventet": "Operatør sjekker ny informasjon, beroliger og avslutter.",
    },
    {
        "rolle": "AMK",
        "spiller": "Johannes",
        "mål": "Samvirke om skadde, fastklemt/innesperret person og ambulanseressurs.",
        "vet": "Får første info via trippel. Trenger antall skadde, bevissthet, fastklemt og trygg adkomst.",
        "vet_ikke": "Skadegrad før første ressurs/ambulanse er fremme.",
        "ikke_impro": "Ikke overta brannfaglige vurderinger.",
        "forventet": "110 deler bekreftet/ubekreftet informasjon tydelig og oppdaterer ved vindusmelding.",
    },
    {
        "rolle": "Politi",
        "spiller": "Ruben",
        "mål": "Samvirke om trafikk, sikring og avklaring av militær kolonne som kontekst.",
        "vet": "Trafikkulykke på E18. Militær kolonne kan være involvert. Politiet eier trafikk/sikring.",
        "vet_ikke": "Detaljer om last før kolonneleder er kontaktet.",
        "ikke_impro": "Ikke legg inn sikkerhetspolitisk hendelse, trussel eller sabotasje.",
        "forventet": "110 gir god status og skiller kjent/ukjent informasjon.",
    },
    {
        "rolle": "VTS",
        "spiller": "Benjamin",
        "mål": "Støtte trafikkavvikling og fremkommelighet.",
        "vet": "E18 retning Oslo får rask kø. Kan bidra med stenging/varsling og trafikkstatus.",
        "vet_ikke": "Skadegrad og brannomfang.",
        "ikke_impro": "Ikke skap stor omkjøringskrise.",
        "forventet": "110 bruker VTS der trafikk/fremkommelighet er relevant og holder status kort.",
    },
    {
        "rolle": "B43-09 / UL",
        "spiller": "Mats",
        "mål": "Gi underveismelding, be om oppdatert info og levere vindusmelding.",
        "vet": "Ressurs er underveis og fremme ca. 22 minutter inn i øvelsen.",
        "vet_ikke": "Full situasjon før fremkomst.",
        "ikke_impro": "Ikke endre skadeomfang uten øvingsleder.",
        "forventet": "110 gir oppdatert situasjonsbilde og støtter uten å overta taktikk.",
    },
    {
        "rolle": "IL brann",
        "spiller": "Mats",
        "mål": "Starte driftsfase og be 110 om samlet status.",
        "vet": "Brann/røyk er avgrenset. Trafikk og pasienthåndtering pågår.",
        "vet_ikke": "Alt som er loggført på sentralen.",
        "ikke_impro": "Ikke eskaler til stor hendelse.",
        "forventet": "VL/OP kan oppsummere status strukturert.",
    },
]


REQUIRED = {
    "01_Ovelsesbeskrivelse_B43_Farrisbrua.docx": ["Farrisbrua", "B43", "3 operatører", "1 vaktleder", "Spilløvelse", "Totalforsvar"],
    "02_Ressursplan_B43_Farrisbrua.docx": ["B43-01", "AMK", "Politi", "VTS", "Militær kolonneleder", "Rune", "Siw", "Tom Christian"],
    "03_Innringerplan_og_spillkort_B43_Farrisbrua.docx": ["Første melder", "Andre melder", "Tredje melder", "Trippelvarsling", "Vibecke", "Benjamin", "Johannes", "Ruben"],
    "04_Dreiebok_B43_Farrisbrua.docx": ["00:00", "22:00", "85:00", "Vindusmelding", "Driftsfase"],
    "05_Sikkerhetsinstruks_B43_Farrisbrua.docx": ["STOPP ØVELSE B43", "reell hendelse", "telefonbruk", "radiobruk"],
    "06_Evalueringsopplegg_ODCR_AAR_B43_Farrisbrua.docx": ["ODCR", "AAR", "Hva forventet vi", "rollefordeling", "situasjonsforståelse"],
    "08_Bildepakke_og_situasjonskort_B43_Farrisbrua.docx": ["Situasjonsbilde", "Bilde 1", "Bilde 4", "Farrisbrua"],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header=True):
    table.style = "Table Grid"
    table.autofit = True
    if header and table.rows:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "203864")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.size = Pt(8.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "203864")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table, header=True)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(str(item))


def base_doc(title, subtitle=None, landscape=False):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        sec.left_margin = Cm(1.25)
        sec.right_margin = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(32, 56, 100)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(160, 48, 42)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(32, 56, 100)
    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(90, 90, 90)

    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    values = [
        ["Sentral", SCENARIO["sentral"], "Sted", SCENARIO["kurssted"]],
        ["Scenario", SCENARIO["sted"], "Versjon", SCENARIO["dato"]],
    ]
    for ridx, row in enumerate(values):
        for cidx, val in enumerate(row):
            set_cell_text(meta.rows[ridx].cells[cidx], val, bold=(cidx % 2 == 0))
            if cidx % 2 == 0:
                set_cell_shading(meta.rows[ridx].cells[cidx], "E9EEF7")
    doc.add_paragraph()
    return doc


def is_empty_paragraph_element(element):
    if element.tag != qn("w:p"):
        return False
    for text in element.iter(qn("w:t")):
        if text.text and text.text.strip():
            return False
    for tag in ("w:drawing", "w:pict", "w:br"):
        if list(element.iter(qn(tag))):
            return False
    return True


def remove_trailing_empty_paragraphs(doc):
    body = doc._body._element
    children = list(body)
    idx = len(children) - 1
    if idx >= 0 and children[idx].tag == qn("w:sectPr"):
        idx -= 1
    while idx >= 0 and is_empty_paragraph_element(children[idx]):
        body.remove(children[idx])
        idx -= 1


def save_doc(doc, filename):
    path = OUT / filename
    remove_trailing_empty_paragraphs(doc)
    doc.save(path)
    return path


def build_images():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    expected = [IMG_DIR / f"bilde_{idx}_farrisbrua.png" for idx in range(1, 5)]
    if all(path.exists() for path in expected):
        return expected

    font = ImageFont.load_default()
    title_font = ImageFont.load_default()
    images = []

    def draw_label(draw, xy, text, fill=(255, 255, 255), outline=(32, 56, 100)):
        x, y = xy
        bbox = draw.textbbox((x, y), text, font=font)
        pad = 6
        rect = (bbox[0]-pad, bbox[1]-pad, bbox[2]+pad, bbox[3]+pad)
        draw.rounded_rectangle(rect, radius=6, fill=fill, outline=outline, width=2)
        draw.text((x, y), text, font=font, fill=(20, 25, 30))

    for idx, title in enumerate([
        "Bilde 1 - Oversikt Farrisbrua, initial melding",
        "Bilde 2 - Nærbilde personbil og militært kjøretøy",
        "Bilde 3 - Trafikk og fremkommelighet",
        "Bilde 4 - Første ressurs fremme / driftsfase",
    ], start=1):
        img = Image.new("RGB", (1600, 900), (230, 240, 246))
        d = ImageDraw.Draw(img)
        d.rectangle((0, 0, 1600, 95), fill=(32, 56, 100))
        d.text((40, 34), title, font=title_font, fill=(255, 255, 255))
        d.rectangle((0, 690, 1600, 900), fill=(83, 130, 158))
        d.text((50, 805), "Farrisvannet / fiktiv skisse", font=font, fill=(255, 255, 255))
        d.rounded_rectangle((120, 260, 1480, 610), radius=28, fill=(85, 87, 91), outline=(55, 55, 60), width=4)
        for y in [342, 432, 522]:
            d.line((150, y, 1450, y), fill=(240, 240, 240), width=4)
        d.line((220, 392, 1380, 392), fill=(255, 220, 90), width=3)
        d.line((220, 482, 1380, 482), fill=(255, 220, 90), width=3)
        d.polygon([(1380, 315), (1450, 335), (1380, 355)], fill=(255, 255, 255))
        d.text((1240, 290), "Retning Oslo/Sandefjord", font=font, fill=(255, 255, 255))

        if idx == 1:
            d.rectangle((740, 390, 835, 440), fill=(185, 30, 30), outline=(20, 20, 20), width=3)
            d.rectangle((860, 374, 1010, 454), fill=(80, 95, 80), outline=(20, 20, 20), width=3)
            for sx, sy, r in [(760, 370, 16), (790, 350, 24), (825, 330, 32)]:
                d.ellipse((sx-r, sy-r, sx+r, sy+r), fill=(80, 80, 80))
            draw_label(d, (600, 220), "Første melder 30-50 m bak")
            draw_label(d, (900, 250), "Militært logistikkjøretøy")
            draw_label(d, (700, 470), "Røyk fra personbil")
        elif idx == 2:
            d.rectangle((620, 380, 760, 455), fill=(185, 30, 30), outline=(20, 20, 20), width=4)
            d.rectangle((790, 360, 1010, 480), fill=(80, 95, 80), outline=(20, 20, 20), width=4)
            d.line((760, 415, 790, 410), fill=(230, 230, 230), width=5)
            for sx, sy, r in [(640, 350, 22), (690, 325, 30), (735, 300, 38)]:
                d.ellipse((sx-r, sy-r, sx+r, sy+r), fill=(75, 75, 75))
            draw_label(d, (500, 240), "1 person innestengt/fastklemt")
            draw_label(d, (820, 500), "Diesellukt, lekkasje ikke bekreftet")
            draw_label(d, (1000, 325), "Kolonne som kontekst")
        elif idx == 3:
            for x in range(260, 700, 85):
                d.rectangle((x, 390, x+58, 428), fill=(210, 210, 210), outline=(20, 20, 20), width=2)
            d.rectangle((760, 390, 855, 440), fill=(185, 30, 30), outline=(20, 20, 20), width=3)
            d.rectangle((880, 374, 1030, 454), fill=(80, 95, 80), outline=(20, 20, 20), width=3)
            d.rectangle((1070, 360, 1160, 480), fill=(244, 196, 48), outline=(20, 20, 20), width=3)
            draw_label(d, (240, 250), "Kø bygger seg opp")
            draw_label(d, (1040, 500), "VTS: stenging/fremkommelighet")
            draw_label(d, (720, 470), "Ulykkespunkt")
        else:
            d.rectangle((680, 390, 775, 440), fill=(185, 30, 30), outline=(20, 20, 20), width=3)
            d.rectangle((805, 374, 955, 454), fill=(80, 95, 80), outline=(20, 20, 20), width=3)
            d.rectangle((1030, 360, 1160, 455), fill=(190, 20, 20), outline=(20, 20, 20), width=3)
            d.text((1050, 392), "B43-01", font=font, fill=(255, 255, 255))
            d.rectangle((1180, 360, 1300, 455), fill=(230, 230, 255), outline=(20, 20, 20), width=3)
            d.text((1200, 392), "AMK", font=font, fill=(20, 20, 20))
            draw_label(d, (500, 245), "Vindusmelding: ikke åpen flamme")
            draw_label(d, (980, 500), "Trygg arbeidsplass etableres")
            draw_label(d, (700, 470), "Driftsfase")

        d.text((40, 858), "ØVELSE - fiktiv situasjonsskisse for B43. Ikke kartgrunnlag.", font=font, fill=(255, 255, 255))
        path = IMG_DIR / f"bilde_{idx}_farrisbrua.png"
        img.save(path)
        images.append(path)
    return images


def build_ovelsesbeskrivelse():
    doc = base_doc("Øvelsesbeskrivelse", SCENARIO["subtitle"])
    doc.add_heading("1. Formål og hensikt", level=1)
    doc.add_paragraph(
        "Hensikten med øvelsen er å utvikle vaktlagets evne til å håndtere en realistisk "
        "trafikk- og brannhendelse på E18 Farrisbrua, med totalforsvarsrelevant kontekst, "
        "gjennom strukturert meldingsmottak, tidlig utalarmering, samvirke og overgang til driftsfase."
    )
    doc.add_heading("2. Målgruppe og rammer", level=1)
    add_table(doc, ["Tema", "Innhold"], [
        ["Målgruppe", "Ett vaktlag på B43: 3 operatører og 1 vaktleder."],
        ["Øvingsmiljø", "Kontrollrom i Stavern med seks OP-plasser. Fire plasser benyttes aktivt."],
        ["Øvingsform", "Spilløvelse med innringere, nødetatsmotspill, VTS og spilte brannressurser."],
        ["Varighet", "Ca. 90 minutter action time, der akuttfasen forventes å være ca. 25-35 minutter."],
        ["Avgrensning", "Systembruk/LEO og lokale rutiner evalueres ikke utover det som er nødvendig for å gjennomføre øvelsen."],
        ["Totalforsvar", "Militær kolonne er scenario-kontekst. Øvelsen evaluerer sivilt nødetatssamvirke, ikke militære prosedyrer."],
    ])
    doc.add_heading("3. Scenario", level=1)
    doc.add_paragraph(
        "En militær logistikkolonne passerer gjennom Larvik-området på E18. På Farrisbrua, retning Oslo/Sandefjord, "
        "skjer en trafikkulykke mellom et sivilt kjøretøy og et militært logistikkjøretøy. Første melding er uklar: "
        "det meldes om røyk, mulig brann, skadde personer, trafikkstans og usikkerhet rundt drivstoff/lekkasje. "
        "Hendelsen utvikler seg naturlig fra akuttfase til driftsfase når første brannressurs kommer frem."
    )
    doc.add_heading("4. Hovedmål og delmål", level=1)
    add_table(doc, ["ID", "Delmål", "Indikator for måloppnåelse"], [[o["id"], o["mål"], o["indikator"]] for o in OBJECTIVES])
    doc.add_heading("5. Praktisk gjennomføring", level=1)
    add_bullets(doc, [
        "Øvingsleder starter og stopper spillet.",
        "Første innringer skal gi grunnlag for trippelvarsling og utalarmering.",
        "Innringer 2 og 3 er duplikat-/avstandsmeldinger i initialfasen og skal avsluttes kort når ny kritisk informasjon er avklart.",
        "Brannressurser spilles via terminal/markør som 01/09, UL og IL.",
        "AMK, politi og VTS spilles via egne telefoner/ICCS der det er mulig.",
        "Øvelsen fryses rundt 85 minutter for sluttstatus og overgang til AAR.",
    ])
    doc.add_heading("6. Hva evalueres ikke", level=1)
    add_bullets(doc, [
        "Detaljert lokal LEO-kompetanse.",
        "Kjennskap til reelle lokale ressursnavn og reelle lokale objekter.",
        "Militære prosedyrer, taktikk eller intern militær kommando.",
        "Kommunal kriseledelse eller større stabsetablering.",
        "Langvarig mediehåndtering.",
    ])
    doc.add_heading("7. Malgrunnlag", level=1)
    add_bullets(doc, [
        "MAL øvelsesbeskrivelse..doc",
        "ØvelsesMal.docx",
        "DSB grunnbok for planlegging, gjennomføring og evaluering av øvelser",
        "Metodehefte evaluering av øvelser",
    ])
    return save_doc(doc, "01_Ovelsesbeskrivelse_B43_Farrisbrua.docx")


def build_ressursplan():
    doc = base_doc("Ressursplan og rolleoversikt", SCENARIO["subtitle"])
    doc.add_heading("1. Øvende funksjoner", level=1)
    add_table(doc, ["Rolle", "Status", "Primæroppgave"], RESOURCES[:4])
    doc.add_heading("2. Motspill og spilte ressurser", level=1)
    add_table(doc, ["Ressurs/rolle", "Type", "Funksjon i øvelsen"], RESOURCES[4:])
    doc.add_heading("3. Personfordeling for gruppe 3 og spillstab", level=1)
    doc.add_paragraph(
        "Gruppe 3 eier øvelsen. To fra gruppe 3 sitter i kontrollrommet som observatører, "
        "og to fra gruppe 3 inngår i spillstab. Gruppe 1 støtter med innringere og motspill."
    )
    add_table(doc, ["Person", "Gruppe", "Plassering", "Rolle", "Primæransvar"], ROLE_ALLOCATION)
    doc.add_heading("4. Kanaler og praktisk bruk", level=1)
    add_table(doc, ["Kanal", "Bruk", "Merknad"], [
        ["Nødtelefon", "Tre innringere i initialfasen", "Innringer 1 skal trippelvarsles; innringer 2 og 3 avsluttes kort."],
        ["ICCS/trippel", "AMK og politi", "Brukes som tilgjengelig i kontrollrommet."],
        ["Telefon", "Politi, AMK, HRS, VTS og andre støtteaktører", "Egne øvingstelefoner. Reelle eksterne nummer skal ikke brukes."],
        ["Terminal/Nødnett", "01/09, UL og IL", "Markør spiller meldinger etter dreiebok."],
        ["Bilder", "Situasjonsforståelse", "Gis i henhold til bildekort, ikke samlet ved start."],
    ])
    doc.add_heading("5. Ressursutvikling i spillet", level=1)
    add_table(doc, ["Tid", "Ressurs", "Status", "Forventet 110-oppgave"], [
        ["00:00-02:00", "B43-01/B43-09", "Utalarmeres", "Rask respons på første melding."],
        ["08:00-15:00", "B43-01/B43-09", "Underveis", "Oppdatere ressurser med posisjon, skadde, røyk/brann og trafikk."],
        ["22:00", "B43-09/UL", "Fremme", "Motta vindusmelding og oppdatere AMK, politi og VTS."],
        ["35:00", "IL brann", "Fremme/driftsfase", "Gi samlet status og avklare videre oppfølging."],
        ["45:00-75:00", "VTS/politi/AMK/IL", "Driftsfase", "Korte statusoppdateringer og avklaringer."],
    ])
    doc.add_heading("6. Foreslått rollefordeling for vaktlaget", level=1)
    doc.add_paragraph("Dette er ikke en fasit, men en forventet god struktur observatørene kan se etter.")
    add_table(doc, ["Funksjon", "Primærrolle", "Sekundærrolle"], [
        ["Meldingsmottak", "OP 1", "OP 3 kan støtte/logge."],
        ["Utalarmering og samvirke", "OP 2", "VL prioriterer ved samtidighet."],
        ["Logg og status", "OP 3", "VL kvalitetssikrer hovedpunkter."],
        ["Ledelse og prioritering", "VL", "Alle bidrar med kort status ved behov."],
    ])
    return save_doc(doc, "02_Ressursplan_B43_Farrisbrua.docx")


def build_innringer_spillkort(images):
    doc = base_doc("Innringerplan og spillkort", SCENARIO["subtitle"])
    doc.add_heading("1. Innringerplan", level=1)
    add_table(doc, ["Nr.", "Tid", "Spilles av", "Rolle", "Kanal", "Hovedfunksjon"], [
        [c["nr"], c["tid"], c["spilles_av"], c["rolle"], c["kanal"], c["forventet"]] for c in CALLERS
    ])
    for c in CALLERS:
        doc.add_heading(f"Innringer {c['nr']}: {c['navn']}", level=2)
        add_table(doc, ["Felt", "Innhold"], [
            ["Spilles av", c["spilles_av"]],
            ["Rolle", c["rolle"]],
            ["Tidspunkt", c["tid"]],
            ["Uttrykk", c["uttrykk"]],
            ["Første melding", c["førstemelding"]],
            ["Gir frivillig", "\n".join(c["frivillig"])],
            ["Gir ved gode spørsmål", "\n".join(c["ved_spørsmål"])],
            ["Usikker/feil informasjon", c["usikkerhet"]],
            ["Avslutning", c["avslutning"]],
            ["Forventet 110-håndtering", c["forventet"]],
        ])

    doc.add_heading("2. Felles bildegrunnlag for spillkort", level=1)
    doc.add_paragraph(
        "Disse bildene er felles visuell referanse for innringere, motspillere og spillstab. "
        "De skal bidra til at alle beskriver samme hendelse. Bildene skal ikke vises samlet til de øvende ved start."
    )
    add_table(doc, ["Bilde", "Brukes av", "Tidspunkt", "Hva bildet skal støtte"], [
        ["Bilde 1", "Melder 1 / spillstab", "0-8 min", "Initial oversikt, røyk, kjøretøyplassering og trafikkstans."],
        ["Bilde 2", "Melder 1 / UL / AMK-politi ved avklaring", "15-22 min", "Nærmere skadepunkt, røykutvikling og personbil mot militært kjøretøy."],
        ["Bilde 3", "VTS / politi / spillstab", "35-45 min", "Kø, blokkert kjørebane og fremkommelighet."],
        ["Bilde 4", "UL / IL / AMK / politi", "22-35 min eller driftsfase", "Første ressurser fremme og arbeid med uthenting/frigjøring."],
    ])
    doc.add_page_break()
    for i, img in enumerate(images, 1):
        if i > 1:
            doc.add_page_break()
        doc.add_heading(f"Bilde {i} - felles referanse", level=2)
        doc.add_picture(str(img), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. Spillkort for øvrige motspillere", level=1)
    for s in SPILLCARDS:
        doc.add_heading(s["rolle"], level=2)
        add_table(doc, ["Felt", "Innhold"], [
            ["Spilles av", s["spiller"]],
            ["Mål med rollen", s["mål"]],
            ["Rollen vet", s["vet"]],
            ["Rollen vet ikke", s["vet_ikke"]],
            ["Skal ikke improvisere på", s["ikke_impro"]],
            ["Forventet respons fra de øvende", s["forventet"]],
        ])
    return save_doc(doc, "03_Innringerplan_og_spillkort_B43_Farrisbrua.docx")


def build_dreiebok():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.0)
    sec.right_margin = Cm(1.0)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(32, 56, 100)
    p = doc.add_paragraph()
    r = p.add_run("Dreiebok for spillstab - Farrisbrua B43. Skal ikke deles med de øvende.")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(32, 56, 100)
    p.paragraph_format.space_after = Pt(0)
    add_table(doc, ["Fase", "Tid", "Hensikt"], [
        ["1. Nødtelefon og første tiltak", "00:00-08:00", "Meldingsmottak, utalarmering og trippelvarsling."],
        ["2. Fremkjøring og felles situasjonsforståelse", "08:00-22:00", "Avklaringer, samvirke og støtte til ressurser underveis."],
        ["3. Første ressurs fremme", "22:00-35:00", "Vindusmelding og overgang til driftsfase."],
        ["4. Driftsfase", "35:00-75:00", "Korte realistiske innspill og statusavklaringer."],
        ["5. Avslutning", "75:00-90:00", "Sluttstatus og overgang til AAR."],
    ])
    p = doc.add_paragraph()
    r = p.add_run("Detaljert dreiebok")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(32, 56, 100)
    p.paragraph_format.space_after = Pt(0)
    add_table(doc,
        ["Tid", "Til", "Fra", "Hendelse/episode", "Mål", "Forventet handling", "Merknad evaluering"],
        TIMELINE
    )
    return save_doc(doc, "04_Dreiebok_B43_Farrisbrua.docx")


def build_sikkerhet():
    doc = base_doc("Sikkerhetsinstruks og spilleregler", SCENARIO["subtitle"])
    doc.add_heading("1. Formål", level=1)
    doc.add_paragraph(
        "Instruksen skal sikre at øvelsen gjennomføres trygt, tydelig avgrenset fra reell beredskap, "
        "og uten fare for misforståelse mellom spill og virkelighet."
    )
    doc.add_heading("2. Stoppregel", level=1)
    doc.add_paragraph("Stoppkode: STOPP ØVELSE B43.")
    add_bullets(doc, [
        "Alle i øvingsledelse og spillstab kan melde behov for stopp.",
        "Øvingsleder beslutter stopp eller pause.",
        "Ved reell hendelse stoppes øvelsen umiddelbart og kontrollrommet frigjøres.",
        "Alle telefoner/radioer skal tydelig merkes eller brukes som øvingskanaler.",
    ])
    doc.add_heading("3. Risiko og tiltak", level=1)
    add_table(doc, ["Risiko", "Tiltak", "Ansvar"], [
        ["Forveksling med reell hendelse", "Alle spillmeldinger starter ved behov med 'ØVELSE ØVELSE ØVELSE'. Reelle eksterne nummer brukes ikke.", "Øvingsleder"],
        ["Reell hendelse oppstår under øvelse", "Stoppkode benyttes. Øvingsleder avbryter og frigjør personell/utstyr.", "Øvingsleder/VL"],
        ["For høy belastning eller uklar læring", "Øvelsen skal ikke eskaleres uten beslutning fra øvingsleder.", "Spillstab"],
        ["Uønsket bruk av bilder", "Bildepakken er fiktiv og merket ØVELSE. Ikke del utenfor kurs-/øvingsformål.", "Dokumentansvarlig"],
        ["Teknisk svikt", "Fallback er muntlig innspill fra spillstab etter dreiebok.", "Teknisk ansvarlig"],
    ])
    doc.add_heading("4. Spilleregler", level=1)
    add_bullets(doc, [
        "De øvende skal håndtere innspill som om de var reelle, innenfor øvingsmiljøets rammer.",
        "Motspillere skal holde seg til spillkort og dreiebok.",
        "Improvisasjon skal ikke endre skadeomfang, farlig last, antall skadde eller større konsekvenser.",
        "Forsvaret/militær kolonne er kontekst. Øvelsen skal ikke utvikles til sikkerhetspolitisk hendelse.",
        "Øvelsen avsluttes med kort sluttstatus fra vaktlaget før AAR.",
    ])
    doc.add_heading("5. Telefonbruk og radiobruk", level=1)
    add_bullets(doc, [
        "telefonbruk: bare avtalte øvingstelefoner og interne øvingsnummer skal benyttes. Reelle nød- eller etatsnummer skal ikke ringes.",
        "Radiobruk: terminaler/Nødnett brukes bare i avtalte øvingstalegrupper eller som simulert samband etter dreiebok.",
        "Alle muntlige avklaringer som ikke kan kjøres teknisk, kan gis som spillteknisk innspill fra øvingsleder.",
    ])
    doc.add_heading("6. Teknisk klargjøring", level=1)
    add_table(doc, ["Kontrollpunkt", "Utføres av", "Status"], [
        ["Telefoner til innringere og motspill testet", "Spillstab", "Før start"],
        ["ICCS/trippel/SAR-funksjon avklart", "Teknisk ansvarlig", "Før start"],
        ["Terminaler for 01/09/UL/IL avklart", "Spillstab", "Før start"],
        ["Bilder tilgjengelig i riktig rekkefølge", "Øvingsleder", "Før start"],
        ["AAR-rom/whiteboard klart", "Evalueringsansvarlig", "Før start"],
    ])
    return save_doc(doc, "05_Sikkerhetsinstruks_B43_Farrisbrua.docx")


def build_evaluering():
    doc = base_doc("Evalueringsopplegg ODCR og AAR", SCENARIO["subtitle"])
    doc.add_heading("1. Evalueringsdesign", level=1)
    doc.add_paragraph(
        "Evalueringen gjennomføres med observasjon under øvelsen og AAR umiddelbart etterpå. "
        "Fokus er måloppnåelse, læring og forbedring, ikke å ta enkeltpersoner."
    )
    doc.add_heading("2. Observasjonspunkter", level=1)
    add_table(doc, ["Mål", "Hva observeres", "Tegn på god måloppnåelse"], [
        ["D1", "Første meldingsmottak og utalarmering", "Posisjon, skade, brann/røyk og risiko avklares raskt. Tiltak startes tidlig."],
        ["D2", "Felles situasjonsforståelse", "VL og OP-er deler korte, presise statusoppdateringer."],
        ["D3", "Samvirke", "AMK, politi og VTS får relevant informasjon og oppdateres ved endring."],
        ["D4", "Rollefordeling og driftsfase", "Oppgaver fordeles tydelig, og sentralen skifter arbeidsform når første ressurs er fremme."],
        ["D5", "Sluttstatus", "Vaktlaget kan oppsummere hendelse, tiltak, varsling og gjenstående oppfølging."],
    ])
    doc.add_paragraph(
        "Observatørene skal særlig merke seg når rollefordeling etableres, om rollefordelingen er tydelig for hele vaktlaget, "
        "og om den justeres hensiktsmessig ved overgang til driftsfase."
    )
    doc.add_heading("3. ODCR-skjema", level=1)
    add_table(doc, ["Nr.", "Observasjon", "Refleksjon", "Konklusjon", "Anbefaling"], [
        ["1", "", "", "", ""],
        ["2", "", "", "", ""],
        ["3", "", "", "", ""],
        ["4", "", "", "", ""],
        ["5", "", "", "", ""],
    ])
    doc.add_heading("4. AAR-gjennomføring", level=1)
    doc.add_paragraph("AAR gjennomføres direkte etter øvelsen. Foreslått tidsbruk: 25-35 minutter.")
    add_table(doc, ["Steg", "Spørsmål", "Hensikt"], [
        ["1", "Hva forventet vi ville skje?", "Koble tilbake til mål og før-brief."],
        ["2", "Hva skjedde?", "Etablere felles faktabilde."],
        ["3", "Hvorfor ble det slik?", "Forstå valg, prioriteringer og rammer."],
        ["4", "Hva gikk bra, og hvorfor?", "Identifisere bevaringspunkter."],
        ["5", "Hva kan forbedres, og hvordan?", "Identifisere konkrete læringspunkter."],
    ])
    doc.add_heading("5. Fortsette, slutte, begynne", level=1)
    add_table(doc, ["Fortsette å gjøre", "Slutte å gjøre", "Begynne å gjøre"], [
        ["", "", ""],
        ["", "", ""],
        ["", "", ""],
    ])
    return save_doc(doc, "06_Evalueringsopplegg_ODCR_AAR_B43_Farrisbrua.docx")


def build_bilde_doc(images):
    doc = base_doc("Bildepakke og situasjonskort", SCENARIO["subtitle"])
    doc.add_heading("1. Bruk av bilder", level=1)
    doc.add_paragraph(
        "Bildene er fiktive, fotorealistiske situasjonsbilder som skal støtte situasjonsforståelse. Hvert Situasjonsbilde skal gis på riktig tidspunkt i dreieboken og ikke deles samlet med de øvende ved start. Samme bildegrunnlag brukes i spillkortene, slik at innringere, motspillere og spillstab beskriver samme hendelse."
    )
    rows = [
        ["Bilde 1", "10-15 min", "Melder/politi", "Oversikt over ulykkespunkt og retning. Brukes hvis de øvende trenger visuell støtte."],
        ["Bilde 2", "15-22 min", "Melder/UL", "Nærbilde av personbil og militært kjøretøy. Støtter brann/røyk og skadeforståelse."],
        ["Bilde 3", "35-45 min", "VTS", "Viser kø og fremkommelighetsproblem. Støtter trafikkavklaring."],
        ["Bilde 4", "22-35 min eller driftsfase", "UL/IL", "Viser første ressurs fremme og overgang til driftsfase."],
    ]
    add_table(doc, ["Bilde", "Tidspunkt", "Gis av", "Hensikt"], rows)
    for i, img in enumerate(images, 1):
        doc.add_page_break()
        doc.add_heading(f"Bilde {i}", level=2)
        doc.add_picture(str(img), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return save_doc(doc, "08_Bildepakke_og_situasjonskort_B43_Farrisbrua.docx")


def extract_docx_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def verify_doc(path):
    text = extract_docx_text(path)
    missing = [needle for needle in REQUIRED.get(path.name, []) if needle not in text]
    return {"file": path.name, "ok": not missing, "missing": missing, "chars": len(text)}


def write_qa(results, images):
    qa_lines = [
        "# QA - Farrisbrua B43",
        "",
        f"Generert: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "## Dokumentkontroll",
        "",
    ]
    for r in results:
        status = "OK" if r["ok"] else "MANGLER"
        qa_lines.append(f"- {status}: `{r['file']}` ({r['chars']} tegn)")
        if r["missing"]:
            qa_lines.append(f"  - Mangler: {', '.join(r['missing'])}")
    qa_lines.extend([
        "",
        "## Bildekontroll",
        "",
    ])
    for img in images:
        with Image.open(img) as im:
            qa_lines.append(f"- OK: `{img.relative_to(OUT)}` {im.size[0]}x{im.size[1]} px")
    qa_lines.extend([
        "",
        "## Scenarioavklaringer valgt i førsteutkast",
        "",
        "- Geografi: E18 Farrisbrua, Larvik, retning Oslo/Sandefjord.",
        "- Innringere: 3 i initialfasen.",
        "- Innringer 1: grunnlag for trippelvarsling og utalarmering.",
        "- Innringer 2 og 3: avstandsmeldinger som avsluttes kort etter sjekk for ny kritisk informasjon.",
        "- Militær kolonne: totalforsvarskontekst, ikke tung motspiller.",
        "- Aktivt motspill: AMK, politi, VTS, 01/09/UL/IL.",
    ])
    (OUT / "00_QA_sjekkliste.md").write_text("\n".join(qa_lines), encoding="utf-8")


def main():
    OUT.mkdir(exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    QA.mkdir(exist_ok=True)

    images = build_images()
    paths = [
        build_ovelsesbeskrivelse(),
        build_ressursplan(),
        build_innringer_spillkort(images),
        build_dreiebok(),
        build_sikkerhet(),
        build_evaluering(),
        build_bilde_doc(images),
    ]
    results = [verify_doc(p) for p in paths]
    write_qa(results, images)
    manifest = {
        "scenario": SCENARIO,
        "documents": [str(p.relative_to(OUT)) for p in paths],
        "images": [str(p.relative_to(OUT)) for p in images],
        "qa": "00_QA_sjekkliste.md",
        "all_ok": all(r["ok"] for r in results),
        "results": results,
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    if not manifest["all_ok"]:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
